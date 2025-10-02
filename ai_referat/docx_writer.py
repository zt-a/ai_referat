import json
import re
from typing import Any, Dict, Optional, Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, RGBColor

from ai_referat.models import Essay


def apply_markdown_formatting(text: str) -> str:
    """Простейшее применение Markdown стилей к тексту (жирный, заголовки)."""
    text = re.sub(r"\*\*(.*?)\*\*", r"**\1**", text)
    text = re.sub(r"### (.*?)\n", r"### \1\n", text)
    text = re.sub(r"## (.*?)\n", r"## \1\n", text)
    text = re.sub(r"# (.*?)\n", r"# \1\n", text)
    return text

def create_docx_file(
    docx_path: str,
    json_data: Optional[Union[Dict[str, Any], Essay]] = None,
    json_path: Optional[str] = None,
    discipline: str = "_______________",
    department: str = "________________________",
    topic_name: str = "____________________________________________",
    author: str = "______________________",
    group: str = "_____________________________",
    checked_by: str = "_________________________",
    year: str = "2024",
    city: str = "Бишкек",
    content_font: str = "Aptos",
    content_size: int = 14
) -> None:
    """
    Создает DOCX файл на основе данных из модели Essay или словаря.
    """
    if json_data is None:
        if json_path is None:
            raise ValueError("Нужно передать либо json_data, либо json_path")
        with open(json_path, "r", encoding="utf-8") as f:
            json_data = json.load(f)

    # Если передана Pydantic модель, преобразуем в dict
    if hasattr(json_data, "dict"):
        data: Dict[str, Any] = json_data.dict()
    else:
        data = json_data  # type: ignore

    doc = Document()

    # --- Шапка ---
    for line in [
        "МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ КР",
        "Государственное учреждение высшего профессионального образования",
        "\n\n\n_________________________________________________",
        "СРС"
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    last_paragraph = doc.paragraphs[-1]
    if last_paragraph.runs:
        run = last_paragraph.runs[0]
        run.font.size = Pt(100)
        run.font.bold = True

    # Метаданные
    for text, alignment in [
        (f"Дисциплина: {discipline}", WD_ALIGN_PARAGRAPH.CENTER),
        (f"Кафедра: {department}", WD_ALIGN_PARAGRAPH.CENTER),
        (f"Тема: {topic_name}", WD_ALIGN_PARAGRAPH.CENTER),
        (f"Выполнил: {author}", WD_ALIGN_PARAGRAPH.RIGHT),
        (f"Группа: {group}", WD_ALIGN_PARAGRAPH.RIGHT),
        (f"Проверил: {checked_by}", WD_ALIGN_PARAGRAPH.RIGHT),
        (f"{city}, {year}", WD_ALIGN_PARAGRAPH.CENTER)
    ]:
        p = doc.add_paragraph(text)
        p.alignment = alignment

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- Заголовок темы ---
    title = doc.add_heading(data.get("topic", "Тема"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- План ---
    plan_heading = doc.add_heading("План", level=1)
    plan_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for chapter in data.get("plan", {}).get("chapters", []):
        chap_paragraph = doc.add_paragraph(chapter.get("title", ""), style="Heading 2")
        chap_paragraph.style.font.size = Pt(content_size)
        chap_paragraph.style.font.name = content_font
        for sub in chapter.get("subchapters", []):
            doc.add_paragraph(sub, style="List Bullet")
            
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- Введение ---
    intro_heading = doc.add_heading("Введение", level=1)
    intro_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro_text: str = data.get("introduction", {}).get("text", "")
    intro_paragraph = doc.add_paragraph(apply_markdown_formatting(intro_text))
    for run in intro_paragraph.runs:
        run.font.size = Pt(content_size)
        run.font.name = content_font

    # --- Главы ---
    for chapter in data.get("chapters", []):
        chap_title: str = chapter.get("title", "")
        chap_text: str = chapter.get("text", "")
        subchapters: list = chapter.get("subchapters", [])

        # Разрыв страницы перед каждой новой главой
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Заголовок главы
        doc.add_paragraph(chap_title, style="Heading 2")

        # Текст главы
        if chap_text:
            doc.add_paragraph(apply_markdown_formatting(chap_text))

        for sub in subchapters:
            sub_title: str = sub.get("title", "") if isinstance(sub, dict) else str(sub)
            sub_text: str = sub.get("text", "") if isinstance(sub, dict) else ""

            # Заголовок подглавы
            doc.add_paragraph(sub_title, style="Heading 3")

            # Текст подглавы с размером шрифта 14pt
            if sub_text:
                para = doc.add_paragraph()
                run = para.add_run(sub_text)
                run.font.size = Pt(14)
                
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- Заключение ---
    concl_heading = doc.add_heading("Заключение", level=1)
    concl_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    concl_text: str = data.get("conclusion", {}).get("text", "")
    concl_paragraph = doc.add_paragraph(apply_markdown_formatting(concl_text))
    for run in concl_paragraph.runs:
        run.font.size = Pt(content_size)
        run.font.name = content_font
        
        
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- Литература ---
    refs_heading = doc.add_heading("Литература", level=1)
    refs_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    refs_items: list = data.get("references", {}).get("items", [])
    refs_text: str = "\n".join(refs_items)
    refs_paragraph = doc.add_paragraph(apply_markdown_formatting(refs_text))
    for run in refs_paragraph.runs:
        run.font.size = Pt(content_size)
        run.font.name = content_font

    # --- Цвет текста ---
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            
    # --- Принудительно установить шрифт для всех параграфов ---
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            style_name = paragraph.style.name.lower()
            
            if "heading" in style_name:
                run.font.size = Pt(max(content_size + 3, 17))  # заголовки — не меньше 17pt
            else:
                run.font.size = Pt(content_size)  # обычный текст — content_size (по умолч. 14pt)


    doc.save(docx_path)
    print(f"Документ успешно создан: {docx_path}")



def create_docx_file_for_json(
    docx_path: str,
    json_data: Optional[Union[Dict[str, Any], Essay]] = None,
    json_path: Optional[str] = None,
    content_font: str = "Aptos",
    content_size: int = 14
) -> None:
    """
    Создает DOCX файл на основе данных из модели Essay или словаря.
    Метаданные (дисциплина, автор, год и т.п.) берутся из ключа "metadata" в JSON.
    """
    if json_data is None:
        if json_path is None:
            raise ValueError("Нужно передать либо json_data, либо json_path")
        with open(json_path, "r", encoding="utf-8") as f:
            json_data = json.load(f)

    if hasattr(json_data, "dict"):
        data: Dict[str, Any] = json_data.dict()
    else:
        data = json_data  # type: ignore

    # --- Извлечение метаданных из JSON ---
    metadata = data.get("metadata", {})
    discipline = metadata.get("discipline", "_______________")
    department = metadata.get("department", "________________________")
    topic_name = metadata.get("topic_name", "____________________________________________")
    author = metadata.get("author", "______________________")
    group = metadata.get("group", "_____________________________")
    checked_by = metadata.get("checked_by", "_________________________")
    year = metadata.get("year", "2024")
    city = metadata.get("city", "Бишкек")

    doc = Document()

    # --- Шапка ---
    for line in [
        "МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ КР",
        "Государственное учреждение высшего профессионального образования",
        "\n\n\n_________________________________________________",
        "СРС"
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    last_paragraph = doc.paragraphs[-1]
    if last_paragraph.runs:
        run = last_paragraph.runs[0]
        run.font.size = Pt(100)
        run.font.bold = True

    # --- Метаданные ---
    for text, alignment in [
        (f"Дисциплина: {discipline}", WD_ALIGN_PARAGRAPH.CENTER),
        (f"Кафедра: {department}", WD_ALIGN_PARAGRAPH.CENTER),
        (f"Тема: {topic_name}", WD_ALIGN_PARAGRAPH.CENTER),
        (f"Выполнил: {author}", WD_ALIGN_PARAGRAPH.RIGHT),
        (f"Группа: {group}", WD_ALIGN_PARAGRAPH.RIGHT),
        (f"Проверил: {checked_by}", WD_ALIGN_PARAGRAPH.RIGHT),
        (f"{city}, {year}", WD_ALIGN_PARAGRAPH.CENTER)
    ]:
        p = doc.add_paragraph(text)
        p.alignment = alignment

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- Заголовок темы ---
    title = doc.add_heading(data.get("topic", "Тема"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- План ---
    plan_heading = doc.add_heading("План", level=1)
    plan_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for chapter in data.get("plan", {}).get("chapters", []):
        chap_paragraph = doc.add_paragraph(chapter.get("title", ""), style="Heading 2")
        chap_paragraph.style.font.size = Pt(content_size)
        chap_paragraph.style.font.name = content_font
        for sub in chapter.get("subchapters", []):
            doc.add_paragraph(sub, style="List Bullet")
            
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- Введение ---
    intro_heading = doc.add_heading("Введение", level=1)
    intro_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro_text: str = data.get("introduction", {}).get("text", "")
    intro_paragraph = doc.add_paragraph(apply_markdown_formatting(intro_text))
    for run in intro_paragraph.runs:
        run.font.size = Pt(content_size)
        run.font.name = content_font

    # --- Главы ---
    for chapter in data.get("chapters", []):
        chap_title: str = chapter.get("title", "")
        chap_text: str = chapter.get("text", "")
        subchapters: list = chapter.get("subchapters", [])

        # Разрыв страницы перед каждой новой главой
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Заголовок главы
        doc.add_paragraph(chap_title, style="Heading 2")

        # Текст главы
        if chap_text:
            doc.add_paragraph(apply_markdown_formatting(chap_text))

        for sub in subchapters:
            sub_title: str = sub.get("title", "") if isinstance(sub, dict) else str(sub)
            sub_text: str = sub.get("text", "") if isinstance(sub, dict) else ""

            # Заголовок подглавы
            doc.add_paragraph(sub_title, style="Heading 3")

            # Текст подглавы с размером шрифта 14pt
            if sub_text:
                para = doc.add_paragraph()
                run = para.add_run(sub_text)
                run.font.size = Pt(14)
                
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- Заключение ---
    concl_heading = doc.add_heading("Заключение", level=1)
    concl_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    concl_text: str = data.get("conclusion", {}).get("text", "")
    concl_paragraph = doc.add_paragraph(apply_markdown_formatting(concl_text))
    for run in concl_paragraph.runs:
        run.font.size = Pt(content_size)
        run.font.name = content_font
        
        
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # --- Литература ---
    refs_heading = doc.add_heading("Литература", level=1)
    refs_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    refs_items: list = data.get("references", {}).get("items", [])
    refs_text: str = "\n".join(refs_items)
    refs_paragraph = doc.add_paragraph(apply_markdown_formatting(refs_text))
    for run in refs_paragraph.runs:
        run.font.size = Pt(content_size)
        run.font.name = content_font

    # --- Цвет текста ---
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            
    # --- Принудительно установить шрифт для всех параграфов ---
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            style_name = paragraph.style.name.lower()
            
            if "heading" in style_name:
                run.font.size = Pt(max(content_size + 3, 17))  # заголовки — не меньше 17pt
            else:
                run.font.size = Pt(content_size)  # обычный текст — content_size (по умолч. 14pt)


    doc.save(docx_path)
    print(f"Документ успешно создан: {docx_path}")
